#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF 文档自动拆分和重命名处理器
支持多种 OCR 引擎：阿里云 OCR、DeepSeek-OCR、百度飞桨 PaddleOCR 云服务
"""

import os
import sys
import fitz  # PyMuPDF
import base64
from pathlib import Path
from typing import List, Tuple, Optional, Dict
from doc_classifier import DocumentClassifier
from logger import Logger

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：未安装 python-docx，无法读取Word文件。请运行: pip install python-docx")


class AliyunOCR:
    """阿里云市场 OCR 客户端"""

    def __init__(self, app_code: str = None):
        """
        初始化阿里云市场 OCR 客户端

        Args:
            app_code: 阿里云市场 AppCode（可从环境变量读取）
        """
        self.app_code = app_code or os.getenv('ALIYUN_OCR_APPCODE')

        if not self.app_code:
            raise ValueError(
                "请设置阿里云市场 AppCode！\n"
                "方法1: 设置环境变量 ALIYUN_OCR_APPCODE\n"
                "方法2: 在初始化时传入 app_code 参数"
            )

        self.api_url = "https://tysbgpu.market.alicloudapi.com/api/predict/ocr_general"
        print("阿里云市场 OCR 客户端初始化成功")

    def get_display_name(self) -> str:
        """获取 OCR 引擎的显示名称"""
        return "阿里云 OCR"

    def recognize_general(self, image_path: str) -> str:
        """
        通用文字识别

        Args:
            image_path: 图片文件路径

        Returns:
            识别的文本内容
        """
        try:
            import requests
            import json

            # 读取图片并转为 base64
            with open(image_path, 'rb') as f:
                image_data = f.read()
                image_base64 = base64.b64encode(image_data).decode('utf-8')

            # 构建请求
            headers = {
                'Authorization': f'APPCODE {self.app_code}',
                'Content-Type': 'application/json; charset=UTF-8'
            }

            body = {
                "image": image_base64,
                "configure": {
                    "min_size": 16,
                    "output_prob": False,
                    "output_keypoints": False
                }
            }

            # 调用 API
            response = requests.post(
                self.api_url,
                headers=headers,
                data=json.dumps(body),
                timeout=30
            )

            # 解析响应
            if response.status_code == 200:
                result = response.json()
                if result.get('success') and result.get('ret'):
                    # 提取所有文本行
                    text_lines = []
                    for item in result['ret']:
                        if 'word' in item:
                            text_lines.append(item['word'])
                    return '\n'.join(text_lines)
                else:
                    print(f"    OCR 识别失败: {result.get('message', '未知错误')}")
                    return ""
            else:
                print(f"    API 调用失败: HTTP {response.status_code}")
                return ""

        except ImportError:
            print("错误：requests 库未安装")
            print("请运行: pip install requests")
            return ""
        except Exception as e:
            print(f"    阿里云 OCR 识别失败: {e}")
            return ""


class DeepSeekOCR:
    """DeepSeek-OCR 客户端（通过 SiliconFlow API）"""

    def __init__(self, api_key: str = None):
        """
        初始化 DeepSeek-OCR 客户端

        Args:
            api_key: SiliconFlow API Key（可从环境变量读取）
        """
        self.api_key = api_key or os.getenv('SILICONFLOW_API_KEY')

        if not self.api_key:
            raise ValueError(
                "请设置 SiliconFlow API Key！\n"
                "方法1: 设置环境变量 SILICONFLOW_API_KEY\n"
                "方法2: 在初始化时传入 api_key 参数"
            )

        self.api_url = "https://api.siliconflow.cn/v1/chat/completions"
        self.model = "PaddlePaddle/PaddleOCR-VL-1.5"
        print("PaddleOCR-VL-1.5 客户端初始化成功")

    def get_display_name(self) -> str:
        """获取 OCR 引擎的显示名称"""
        return "PaddleOCR-VL-1.5"

    def recognize_general(self, image_path: str) -> str:
        """
        通用文字识别

        Args:
            image_path: 图片文件路径

        Returns:
            识别的文本内容
        """
        try:
            import requests
            import json
            from PIL import Image

            # 读取并压缩图片（降低 token 消耗）
            img = Image.open(image_path)

            # 如果图片太大，压缩到合适的尺寸
            max_size = 1024  # 最大边长
            if max(img.size) > max_size:
                ratio = max_size / max(img.size)
                new_size = tuple(int(dim * ratio) for dim in img.size)
                img = img.resize(new_size, Image.Resampling.LANCZOS)

            # 转换为 RGB（如果是 RGBA）
            if img.mode == 'RGBA':
                img = img.convert('RGB')

            # 保存到临时文件
            import tempfile
            import os
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
                tmp_path = tmp_file.name
                img.save(tmp_path, 'JPEG', quality=85)

            try:
                # 读取压缩后的图片并转为 base64
                with open(tmp_path, 'rb') as f:
                    image_data = f.read()
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
            finally:
                # 删除临时文件
                try:
                    os.unlink(tmp_path)
                except:
                    pass

            # 判断图片格式
            mime_type = 'image/jpeg'

            # 构建请求（OpenAI 兼容格式）
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }

            # 根据模型类型选择合适的 prompt
            if "PaddleOCR" in self.model or "DeepSeek-OCR" in self.model:
                # PaddleOCR-VL 和 DeepSeek-OCR 是专门的 OCR 模型，使用最简单的 prompt
                ocr_prompt = "识别图片中的所有文字"
            else:
                # DeepSeek-VL2 等通用模型，使用详细的 prompt
                ocr_prompt = "请识别图片中的所有文字内容。重要提示：\n1. 优先识别页面顶部的文件编号（格式如：JDB25300-XXXX-01）\n2. 按从上到下、从左到右的顺序输出\n3. 保持原始布局\n4. 不要添加任何解释或说明"

            body = {
                "model": self.model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{image_base64}"
                                }
                            },
                            {
                                "type": "text",
                                "text": ocr_prompt
                            }
                        ]
                    }
                ],
                "stream": False,
                "max_tokens": 2048,  # 降低到 2048，为图片 token 留出空间
                "temperature": 0.1,
                "top_p": 0.7
            }

            # 调用 API
            response = requests.post(
                self.api_url,
                headers=headers,
                data=json.dumps(body),
                timeout=60
            )

            # 解析响应
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    content = result['choices'][0]['message']['content']
                    return content.strip()
                else:
                    print(f"    OCR 识别失败: 响应格式错误")
                    return ""
            else:
                error_msg = response.text
                print(f"    API 调用失败: HTTP {response.status_code}")
                print(f"    错误信息: {error_msg}")
                return ""

        except ImportError as e:
            if 'PIL' in str(e):
                print("错误：Pillow 库未安装")
                print("请运行: pip install Pillow")
            else:
                print("错误：requests 库未安装")
                print("请运行: pip install requests")
            return ""
        except Exception as e:
            print(f"    DeepSeek-OCR 识别失败: {e}")
            return ""


class BaiduPaddleOCR:
    """百度飞桨 PaddleOCR 云服务 API 客户端（AI Studio 免费额度）

    支持的模型：
        - PaddleOCR-VL-1.5（推荐）：94.5% 精度，支持异形框、印章识别
        - PaddleOCR-VL：版式解析 + OCR
        - PP-OCRv5：轻量级文字识别
        - PP-StructureV3：结构化解析

    通过设置不同的 API_URL 选择模型，API 格式通用。
    """

    # 异步 API 公共端点（无需用户专属 URL）
    ASYNC_JOB_URL = "https://paddleocr.aistudio-app.com/api/v2/ocr/jobs"
    DEFAULT_MODEL = "PaddleOCR-VL-1.5"

    def __init__(self, token: str = None, api_url: str = None):
        """
        初始化百度飞桨 PaddleOCR 云服务客户端

        Args:
            token: AI Studio 访问令牌（可从环境变量 BAIDU_PADDLEOCR_TOKEN 读取）
            api_url: 同步 API 地址（可选，每个用户独立的 URL）
                     不设置则自动使用异步 API 公共端点

        获取方式：
            访问 https://aistudio.baidu.com/paddleocr/task
            在 API 调用示例中获取 TOKEN（API_URL 可选）
        """
        self.token = token or os.getenv('BAIDU_PADDLEOCR_TOKEN')
        self.api_url = api_url or os.getenv('BAIDU_PADDLEOCR_URL')
        self.use_async = not self.api_url  # 没有自定义 URL 时使用异步 API

        if not self.token:
            raise ValueError(
                "请设置百度飞桨 PaddleOCR Token！\n"
                "方法1: 设置环境变量 BAIDU_PADDLEOCR_TOKEN\n"
                "方法2: 在初始化时传入 token 参数\n"
                "获取方式: 访问 https://aistudio.baidu.com/paddleocr/task"
            )

        # 从 API URL 识别模型类型
        self.model_name = self._detect_model_name()
        mode = "异步公共API" if self.use_async else "同步专属API"
        print(f"百度飞桨 {self.model_name} 云服务初始化成功（{mode}）")

    def _detect_model_name(self) -> str:
        """从 API URL 路径推断模型名称"""
        if self.use_async:
            return self.DEFAULT_MODEL
        url_lower = self.api_url.lower()
        if 'layout-parsing' in url_lower or 'vl' in url_lower:
            return "PaddleOCR-VL-1.5"
        elif 'general-recognition' in url_lower or 'ocrv5' in url_lower:
            return "PP-OCRv5"
        elif 'structure' in url_lower:
            return "PP-StructureV3"
        return "PaddleOCR"

    def get_display_name(self) -> str:
        """获取 OCR 引擎的显示名称"""
        return f"百度飞桨 {self.model_name}（云服务）"

    def _recognize_sync(self, image_path: str) -> str:
        """同步 API 调用（需要用户专属 URL）"""
        import requests

        with open(image_path, 'rb') as f:
            file_data = base64.b64encode(f.read()).decode('ascii')

        headers = {
            'Authorization': f'token {self.token}',
            'Content-Type': 'application/json'
        }

        payload = {
            'file': file_data,
            'fileType': 1,
            'useDocOrientationClassify': True,   # 自动矫正旋转（0°/90°/180°/270°）
            'useDocUnwarping': True,              # 自动矫正褶皱/倾斜
            'useLayoutDetection': False,          # 关闭版面检测，改用纯 OCR 模式，避免遗漏页眉小字
            'promptLabel': 'ocr',                 # VL 模型 prompt 类型：纯文字识别
            'useChartRecognition': False,          # 不需要图表识别
            'visualize': False,                    # 不需要可视化结果图，减少返回时间
        }

        response = requests.post(self.api_url, json=payload, headers=headers, timeout=60)

        if response.status_code == 200:
            return self._parse_result(response.json())
        else:
            print(f"    API 调用失败: HTTP {response.status_code}")
            print(f"    错误信息: {response.text[:200]}")
            return ""

    def _recognize_async(self, image_path: str) -> str:
        """异步 API 调用（公共端点，无需专属 URL）"""
        import requests
        import json
        import time

        headers = {
            'Authorization': f'bearer {self.token}',
        }

        optional_payload = {
            'useDocOrientationClassify': True,    # 自动矫正旋转（0°/90°/180°/270°）
            'useDocUnwarping': True,               # 自动矫正褶皱/倾斜
            'useLayoutDetection': False,           # 关闭版面检测，改用纯 OCR 模式，避免遗漏页眉小字
            'promptLabel': 'ocr',                  # VL 模型 prompt 类型：纯文字识别
            'useChartRecognition': False,           # 不需要图表识别
            'visualize': False,                     # 不需要可视化结果图，减少返回时间
        }

        # 提交任务（文件上传模式）
        data = {
            'model': self.DEFAULT_MODEL,
            'optionalPayload': json.dumps(optional_payload),
        }

        with open(image_path, 'rb') as f:
            files = {'file': f}
            response = requests.post(
                self.ASYNC_JOB_URL,
                headers=headers,
                data=data,
                files=files,
                timeout=60
            )

        if response.status_code != 200:
            print(f"    异步任务提交失败: HTTP {response.status_code}")
            print(f"    错误信息: {response.text[:200]}")
            return ""

        job_id = response.json().get('data', {}).get('jobId')
        if not job_id:
            print(f"    异步���务提交失败: 未获取到 jobId")
            return ""

        # 轮询等待结果（单页图片通常很快）
        for _ in range(30):  # 最多等待 60 秒
            time.sleep(2)
            result_resp = requests.get(
                f"{self.ASYNC_JOB_URL}/{job_id}",
                headers=headers,
                timeout=30
            )
            if result_resp.status_code != 200:
                continue

            state = result_resp.json().get('data', {}).get('state')
            if state == 'done':
                jsonl_url = result_resp.json()['data'].get('resultUrl', {}).get('jsonUrl')
                if jsonl_url:
                    return self._fetch_async_result(jsonl_url)
                return ""
            elif state == 'failed':
                error_msg = result_resp.json().get('data', {}).get('errorMsg', '��知错误')
                print(f"    异步任务失败: {error_msg}")
                return ""

        print(f"    异步任务超时")
        return ""

    def _fetch_async_result(self, jsonl_url: str) -> str:
        """从异步任务结果 URL 获取识别文本"""
        import requests
        import json
        import time

        # 重试机制，应对临时网络/SSL 错误
        resp = None
        for attempt in range(3):
            try:
                resp = requests.get(jsonl_url, timeout=30)
                resp.raise_for_status()
                break
            except (requests.exceptions.SSLError, requests.exceptions.ConnectionError) as e:
                if attempt < 2:
                    time.sleep(1 * (attempt + 1))
                    continue
                print(f"    结果下载失败（重试 {attempt + 1} 次后）: {e}")
                return ""
        if resp is None:
            return ""

        text_parts = []
        for line in resp.text.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            result = json.loads(line)
            parsed = self._parse_result(result)
            if parsed:
                text_parts.append(parsed)
        return '\n'.join(text_parts)

    def _parse_result(self, result: dict) -> str:
        """解析 API 返回结果，提取文本（原始格式，标准化由上层统一处理）"""
        if 'result' in result:
            res = result['result']
            # PaddleOCR-VL 返回格式：layoutParsingResults[i].markdown.text
            if 'layoutParsingResults' in res:
                text_parts = []
                for page_res in res['layoutParsingResults']:
                    md = page_res.get('markdown', {})
                    text = md.get('text', '')
                    if text:
                        text_parts.append(text)
                return '\n'.join(text_parts)
            # PP-OCRv5 返回格式：ocrResults[i].recText
            elif 'ocrResults' in res:
                text_parts = []
                for ocr_item in res['ocrResults']:
                    if 'recText' in ocr_item:
                        text_parts.append(ocr_item['recText'])
                    elif 'prunedResult' in ocr_item:
                        pruned = ocr_item['prunedResult']
                        if 'rec_texts' in pruned:
                            text_parts.extend(pruned['rec_texts'])
                return '\n'.join(text_parts)
            else:
                print(f"    未知返回格式，可用字段: {list(res.keys())}")
                return str(res)
        elif 'errorCode' in result:
            print(f"    API 错误: {result.get('errorCode')} - {result.get('errorMsg', '')}")
            return ""
        else:
            print(f"    OCR 返回格式异常: {str(result)[:200]}")
            return ""

    def recognize_general(self, image_path: str) -> str:
        """
        通用文字识别

        Args:
            image_path: 图片文件路径

        Returns:
            识别的文本内容
        """
        try:
            if self.use_async:
                return self._recognize_async(image_path)
            else:
                return self._recognize_sync(image_path)

        except ImportError:
            print("错误：requests 库未安装")
            print("请运行: pip install requests")
            return ""
        except Exception as e:
            print(f"    百度飞桨 PaddleOCR 识别失败: {e}")
            return ""


class LocalPaddleOCR:
    """本地 PaddleOCR-VL 客户端（直接加载模型或通过 MLX-VLM 服务）"""

    # 类级别的模型缓存（单例模式）
    _shared_model = None
    _shared_processor = None
    _model_lock = None

    def __init__(self, base_url: str = None, model_path: str = None, use_direct: bool = False):
        """
        初始化本地 PaddleOCR-VL 客户端

        Args:
            base_url: 本地服务地址（默认 http://localhost:8111）
            model_path: 模型路径或模型名称
            use_direct: 是否直接加载模型（不使用 HTTP 服务）
        """
        self.base_url = base_url or os.getenv('LOCAL_PADDLEOCR_URL', 'http://localhost:8111')
        self.model_path = model_path or os.getenv('LOCAL_PADDLEOCR_MODEL', 'mlx-community/PaddleOCR-VL-1.5-bf16')
        self.use_direct = use_direct or os.getenv('LOCAL_PADDLEOCR_DIRECT', '').lower() == 'true'

        # 初始化锁（用于线程安全的模型加载）
        if LocalPaddleOCR._model_lock is None:
            import threading
            LocalPaddleOCR._model_lock = threading.Lock()

        if self.use_direct:
            print(f"本地 PaddleOCR-VL 客户端初始化（直接加载模型）")
            print(f"  模型: {self.model_path}")
        else:
            print(f"本地 PaddleOCR-VL 客户端初始化成功（{self.base_url}）")
            print(f"  模型: {self.model_path}")

    def get_display_name(self) -> str:
        """获取 OCR 引擎的显示名称"""
        return "PaddleOCR-VL-1.5（本地）"

    def recognize_general(self, image_path: str) -> str:
        """
        通用文字识别

        Args:
            image_path: 图片文件路径

        Returns:
            识别的文本内容
        """
        if self.use_direct:
            return self._recognize_direct(image_path)
        else:
            return self._recognize_via_api(image_path)

    def _recognize_direct(self, image_path: str) -> str:
        """直接加载模型进行识别（线程安全）"""
        try:
            from mlx_vlm import load, generate
            from mlx_vlm.prompt_utils import apply_chat_template
            from PIL import Image

            # 使用锁确保模型只加载一次（线程安全）
            with LocalPaddleOCR._model_lock:
                if LocalPaddleOCR._shared_model is None:
                    print(f"  正在加载模型: {self.model_path}")
                    LocalPaddleOCR._shared_model, LocalPaddleOCR._shared_processor = load(self.model_path)
                    print(f"  ✓ 模型加载完成")

            # 加载图片（使用 PIL）
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # 构建 prompt
            prompt = apply_chat_template(
                LocalPaddleOCR._shared_processor,
                config=LocalPaddleOCR._shared_model.config,
                prompt="识别图片中的所有文字",
                num_images=1
            )

            # 生成结果（image 作为关键字参数）
            output = generate(
                LocalPaddleOCR._shared_model,
                LocalPaddleOCR._shared_processor,
                prompt,
                image=image,
                max_tokens=2048,
                temperature=0.0
            )

            # 提取文本内容
            if hasattr(output, 'text'):
                return output.text.strip()
            elif isinstance(output, str):
                return output.strip()
            else:
                return str(output).strip()

        except ImportError as e:
            print(f"    错误: mlx_vlm 或 PIL 库未安装")
            print(f"    请运行: pip install mlx-vlm>=0.3.11 Pillow")
            return ""
        except Exception as e:
            print(f"    本地模型识别失败: {e}")
            import traceback
            traceback.print_exc()
            return ""

    def _recognize_via_api(self, image_path: str) -> str:
        """通过 HTTP API 进行识别"""
        try:
            import requests
            import json
            from PIL import Image

            # 读取并压缩图片
            img = Image.open(image_path)

            # 压缩到合适的尺寸
            max_size = 1024
            if max(img.size) > max_size:
                ratio = max_size / max(img.size)
                new_size = tuple(int(dim * ratio) for dim in img.size)
                img = img.resize(new_size, Image.Resampling.LANCZOS)

            # 转换为 RGB
            if img.mode == 'RGBA':
                img = img.convert('RGB')

            # 保存到临时文件
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
                tmp_path = tmp_file.name
                img.save(tmp_path, 'JPEG', quality=85)

            try:
                # 读取图片并转为 base64
                with open(tmp_path, 'rb') as f:
                    image_data = f.read()
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
            finally:
                # 删除临时文件
                try:
                    os.unlink(tmp_path)
                except:
                    pass

            # 构建请求（MLX-VLM 格式）
            headers = {
                'Content-Type': 'application/json'
            }

            body = {
                "model": self.model_path,
                "input": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": "识别图片中的所有文字"},
                            {"type": "input_image", "image_url": f"data:image/jpeg;base64,{image_base64}"}
                        ]
                    }
                ],
                "max_output_tokens": 2048
            }

            # 调用本地 API
            response = requests.post(
                f"{self.base_url}/responses",
                headers=headers,
                data=json.dumps(body),
                timeout=60
            )

            # 解析响应
            if response.status_code == 200:
                result = response.json()
                if 'output' in result and len(result['output']) > 0:
                    if 'content' in result['output'][0] and len(result['output'][0]['content']) > 0:
                        content = result['output'][0]['content'][0].get('text', '')
                        return content.strip()
                print(f"    OCR 识别失败: 响应格式错误")
                return ""
            else:
                error_msg = response.text
                print(f"    本地 API 调用失败: HTTP {response.status_code}")
                print(f"    错误信息: {error_msg}")
                return ""

        except requests.exceptions.ConnectionError:
            print(f"    连接失败: 无法连接到本地服务 {self.base_url}")
            print(f"    请确保 MLX-VLM 服务正在运行: mlx_vlm.server --port 8111")
            return ""
        except ImportError as e:
            if 'PIL' in str(e):
                print("错误：Pillow 库未安装")
                print("请运行: pip install Pillow")
            else:
                print("错误：requests 库未安装")
                print("请运行: pip install requests")
            return ""
        except Exception as e:
            print(f"    本地 PaddleOCR-VL 识别失败: {e}")
            return ""



class PDFProcessor:
    """PDF 文档处理器（支持多种 OCR 引擎）"""

    def __init__(
        self,
        app_code: str = None,
        api_key: str = None,
        ocr_engine: str = "aliyun",
        text_threshold: int = 50,
        dpi: int = 300,
        max_workers: int = 3
    ):
        """
        初始化处理器

        Args:
            app_code: 阿里云市场 AppCode（使用阿里云 OCR 时需要）
            api_key: SiliconFlow API Key（使用 PaddleOCR 时需要）
            ocr_engine: OCR 引擎选择，可选值：
                - "aliyun": 阿里云 OCR（默认）
                - "deepseek": PaddleOCR-VL-1.5（通过 SiliconFlow）
                - "local": PaddleOCR-VL-1.5（本地 MLX-VLM）
            text_threshold: 文本提取阈值，少于此字符数则使用 OCR
            dpi: OCR 时的图片 DPI
            max_workers: OCR 并发数，默认 3
        """
        self.text_threshold = text_threshold
        self.dpi = dpi
        self.classifier = DocumentClassifier()
        self.ocr = None  # 延迟加载 OCR
        self.ocr_engine = ocr_engine.lower()
        self.app_code = app_code
        self.api_key = api_key
        self.max_workers = max_workers

        # 验证 OCR 引擎选择
        if self.ocr_engine not in ["aliyun", "deepseek", "local", "baidu"]:
            raise ValueError(f"不支持的 OCR 引擎: {ocr_engine}，可选值: aliyun, deepseek, local, baidu")

    def _init_ocr(self):
        """延迟初始化 OCR（避免不需要时加载）"""
        if self.ocr is None:
            if self.ocr_engine == "aliyun":
                print("正在初始化阿里云市场 OCR...")
                self.ocr = AliyunOCR(self.app_code)
            elif self.ocr_engine == "deepseek":
                print("正在初始化 DeepSeek-OCR...")
                self.ocr = DeepSeekOCR(self.api_key)
            elif self.ocr_engine == "baidu":
                print("正在初始化百度飞桨 PaddleOCR 云服务...")
                self.ocr = BaiduPaddleOCR()
            elif self.ocr_engine == "local":
                print("正在初始化本地 PaddleOCR-VL...")
                self.ocr = LocalPaddleOCR()

    def _get_ocr_display_name(self) -> str:
        """
        获取 OCR 引擎的显示名称

        Returns:
            OCR 引擎的显示名称
        """
        # 如果 OCR 已初始化，直接从实例获取
        if self.ocr and hasattr(self.ocr, 'get_display_name'):
            return self.ocr.get_display_name()

        # 如果未初始化，返回默认名称
        if self.ocr_engine == "deepseek":
            return "PaddleOCR-VL-1.5"
        elif self.ocr_engine == "baidu":
            return "百度飞桨 PaddleOCR（云服务）"
        elif self.ocr_engine == "local":
            return "PaddleOCR-VL-1.5（本地）"
        else:
            return "阿里云 OCR"


    def extract_system_name_from_report(self, pdf_dir: Path) -> Optional[str]:
        """
        从测评报告Word文件中提取系统名称

        在PDF所在目录下搜索测评报告Word文件，从首页提取系统名称

        Args:
            pdf_dir: PDF文件所在目录

        Returns:
            系统名称，如"xxx运营系统"，如果未找到则返回 None
        """
        if not DOCX_AVAILABLE:
            print("警告：未安装 python-docx，无法读取测评报告")
            return None

        # 搜索测评报告Word文件（在当前目录、父目录及子目录）
        search_dirs = [pdf_dir, pdf_dir.parent]
        report_file = None

        for search_dir in search_dirs:
            if not search_dir.exists():
                continue

            # 先在当前目录搜索
            for item in search_dir.iterdir():
                # 查找测评报告文件（排除渗透测试报告）
                if item.is_file() and item.suffix in ['.docx', '.doc']:
                    filename = item.name
                    if ('测评报告' in filename or '报告' in filename) and '渗透' not in filename:
                        report_file = item
                        break

                # 如果是目录，且目录名包含"测评报告"，进入子目录搜索
                if item.is_dir() and '测评报告' in item.name:
                    for sub_item in item.iterdir():
                        if sub_item.is_file() and sub_item.suffix in ['.docx', '.doc']:
                            sub_filename = sub_item.name
                            if ('测评报告' in sub_filename or '报告' in sub_filename) and '渗透' not in sub_filename:
                                report_file = sub_item
                                break
                    if report_file:
                        break

            if report_file:
                break

        if not report_file:
            print("提示：未找到测评报告Word文件")
            return None

        print(f"找到测评报告: {report_file.name}")

        try:
            import re
            doc = Document(report_file)

            # 查找首页（前30段）
            for para in doc.paragraphs[:30]:
                text = para.text.strip()

                # 匹配格式1：网络安全等级保护 xxx 等级测评报告（单段）
                # 提取"网络安全等级保护"和"等级测评报告"之间的内容作为系统名称
                pattern1 = r'网络安全等级保护\s*(.+?)\s*等级测评报告'
                match = re.search(pattern1, text)
                if match:
                    system_name = match.group(1).strip()
                    print(f"  从测评报告提取系统名称: {system_name}")
                    return system_name

                # 匹配格式2：xxx等级测评报告（单段，不包含"网络安全等级保护"）
                pattern2 = r'^(.+?)等级测评报告$'
                match = re.search(pattern2, text)
                if match and '网络安全等级保护' not in text:
                    system_name = match.group(1).strip()
                    print(f"  从测评报告提取系统名称: {system_name}")
                    return system_name

            # 如果上面没找到，尝试查找连续的两段
            for i in range(len(doc.paragraphs) - 1):
                para1 = doc.paragraphs[i].text.strip()
                para2 = doc.paragraphs[i + 1].text.strip()

                # 检查是否是"网络安全等级保护" + "xxx等级测评报告"的组合
                if '网络安全等级保护' in para1 and '等级测评报告' in para2:
                    # 从第二段提取系统名称（去掉"等级测评报告"）
                    pattern = r'^(.+?)等级测评报告$'
                    match = re.search(pattern, para2)
                    if match:
                        system_name = match.group(1).strip()
                        print(f"  从测评报告提取系统名称（跨段）: {system_name}")
                        return system_name

            print("  未能从测评报告提取系统名称")
            return None

        except Exception as e:
            print(f"读取测评报告失败: {e}")
            return None

    def read_system_names_from_word(self, pdf_dir: Path) -> Dict[int, str]:
        """
        从项目实施申请单Word文件中读取系统名称

        在PDF所在目录下搜索"项目实施单"文件夹，读取"项目实施申请单"Word文件
        提取"序号"列和"系统名称"列的对应关系

        优先从测评报告中提取系统名称，然后用它更新实施单中的系统名称

        Args:
            pdf_dir: PDF文件所在目录

        Returns:
            字典：{序号: 系统名称}，例如 {1: "门户网站系统", 2: "内部办公系统"}
        """
        if not DOCX_AVAILABLE:
            print("警告：未安装 python-docx，无法读取系统名称")
            return {}

        # 1. 先从测评报告提取系统名称
        report_system_name = self.extract_system_name_from_report(pdf_dir)
        if report_system_name:
            print(f"✓ 从测评报告获取系统名称: {report_system_name}")

        # 2. 搜索"项目实施单"文件夹（先在当前目录，再在父目录）
        impl_folder = None
        search_dirs = [pdf_dir, pdf_dir.parent]  # 当前目录和父目录

        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            for item in search_dir.iterdir():
                if item.is_dir() and "项目实施单" in item.name:
                    impl_folder = item
                    break
            if impl_folder:
                break

        if not impl_folder:
            print(f"错误：未找到包含'项目实施单'的文件夹（已搜索当前目录和父目录）")
            return {}

        print(f"找到项目实施单文件夹: {impl_folder.name}")

        # 3. 搜索"项目实施申请单"Word文件
        word_file = None
        for item in impl_folder.iterdir():
            if item.is_file() and "项目实施申请单" in item.name and item.suffix in ['.docx', '.doc']:
                word_file = item
                break

        if not word_file:
            print(f"错误：未找到'项目实施申请单'Word文件")
            return {}

        print(f"找到Word文件: {word_file.name}")

        try:
            # 4. 读取Word文件
            doc = Document(word_file)
            system_names = {}
            doc_modified = False  # 标记文档是否被修改

            # 遍历所有表格
            for table in doc.tables:
                # 查找表头行
                header_row = None
                seq_col_idx = None
                system_col_idx = None

                for row_idx, row in enumerate(table.rows):
                    cells_text = [cell.text.strip() for cell in row.cells]

                    # 查找"序号"和"系统名称"列
                    for col_idx, text in enumerate(cells_text):
                        if "序号" in text:
                            seq_col_idx = col_idx
                        if "系统名称" in text:
                            system_col_idx = col_idx

                    # 找到"系统名称"列即可（序号列可选）
                    if system_col_idx is not None:
                        header_row = row_idx
                        break

                # 如果找到了表头，读取数据行
                if header_row is not None:
                    auto_seq = 1  # 无序号列时自动编号
                    for row_idx in range(header_row + 1, len(table.rows)):
                        row = table.rows[row_idx]
                        try:
                            system_text = row.cells[system_col_idx].text.strip()

                            if not system_text:
                                continue

                            # 跳过非系统名称的行（如"其他要求"等）
                            skip_keywords = ["其他要求", "项目实施", "备注"]
                            if any(kw in system_text for kw in skip_keywords):
                                continue

                            # 确定序号
                            seq_num = None
                            if seq_col_idx is not None:
                                # 有序号列，使用序号
                                seq_text = row.cells[seq_col_idx].text.strip()
                                if seq_text:
                                    try:
                                        seq_num = int(seq_text)
                                    except ValueError:
                                        pass

                            if seq_num is None:
                                # 无序号列或序号无效，自动编号
                                seq_num = auto_seq
                                auto_seq += 1

                            # 如果从测评报告提取到了系统名称，用它替换实施单中的系统名称
                            original_name = system_text
                            if report_system_name:
                                system_text = report_system_name
                                if original_name != system_text:
                                    # 更新实施单中的单元格
                                    row.cells[system_col_idx].text = system_text
                                    doc_modified = True
                                    print(f"  更新: 序号 {seq_num} → {original_name} 替换为 {system_text}")
                                else:
                                    print(f"  读取: 序号 {seq_num} → {system_text}")
                            else:
                                print(f"  读取: 序号 {seq_num} → {system_text}")

                            system_names[seq_num] = system_text

                        except (ValueError, IndexError):
                            continue

            # 5. 如果文档被修改，保存更新后的实施单
            if doc_modified:
                doc.save(word_file)
                print(f"✓ 已保存更新后的项目实施单")

            if system_names:
                print(f"成功读取 {len(system_names)} 个系统名称")
            else:
                print("错误：未能从Word文件中读取到系统名称")

            return system_names

        except Exception as e:
            print(f"读取Word文件失败: {e}")
            return {}

    def split_pdf(self, pdf_path: str, output_dir: str) -> List[str]:
        """
        拆分 PDF 为单页文件

        Args:
            pdf_path: PDF 文件路径
            output_dir: 输出目录

        Returns:
            拆分后的单页 PDF 文件路径列表
        """
        pdf_path = Path(pdf_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        doc = fitz.open(pdf_path)
        page_files = []

        print(f"\n正在拆分 PDF: {pdf_path.name}")
        print(f"总页数: {len(doc)}")

        for page_num in range(len(doc)):
            # 创建新的 PDF 包含单页
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)

            # 保存单页 PDF
            output_file = output_dir / f"page_{page_num + 1:03d}.pdf"
            new_doc.save(output_file)
            new_doc.close()

            page_files.append(str(output_file))
            print(f"  拆分第 {page_num + 1} 页 → {output_file.name}")

        doc.close()
        return page_files

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        从 PDF 提取文本（混合方案）

        Args:
            pdf_path: PDF 文件路径

        Returns:
            提取的文本内容
        """
        doc = fitz.open(pdf_path)
        page = doc[0]  # 单页 PDF

        # 1. 尝试提取文本层
        text = page.get_text()

        # 2. 如果文本太少，判断为扫描件，使用 OCR
        is_scanned = len(text.strip()) < self.text_threshold
        if is_scanned:
            print(f"    扫描件，OCR识别中...", end='', flush=True)
            text = self._ocr_page(page)

        # 统一文本标准化：去除 HTML 标签等格式差异，确保所有 OCR 引擎输出一致
        text = self._normalize_ocr_text(text)

        # 显示识别结果预览（标准化后的纯文本）
        if is_scanned:
            preview = text[:50].replace('\n', ' ') if text else '(空)'
            print(f" ✓ 识别完成: {preview}...")

        # 补充提取文件编号：OCR 引擎可能丢失页眉小字，用 PyMuPDF 内置 Tesseract 兜底
        if is_scanned and text and not self.classifier.extract_file_code(text):
            header_text = self._ocr_header_for_file_code(page)
            if header_text:
                text = header_text + '\n' + text
                print(f"    [补充] 从页眉提取到文件编号信息")

        doc.close()

        return text.strip()

    def _ocr_header_for_file_code(self, page) -> Optional[str]:
        """
        用 PyMuPDF 内置 Tesseract 对全页做 OCR，专门提取文件编号

        云 OCR（如百度飞桨 VL 模型）经常忽略页眉小字，导致文件编号丢失。
        此方法用本地 Tesseract 做补充识别，仅提取文件编号部分。

        Args:
            page: PyMuPDF 页面对象

        Returns:
            包含文件编号的文本，如果未找到则返回 None
        """
        import re
        try:
            tp = page.get_textpage_ocr(language='chi_sim+eng', dpi=300)
            full_text = page.get_text(textpage=tp)
            if not full_text:
                return None
            # 标准化连字符
            full_text = re.sub(r'[—–−ー一]', '-', full_text)
            # 检查是否包含文件编号
            if self.classifier.extract_file_code(full_text):
                # 只返回前20%文本（包含文件编号的页眉区域），避免干扰正文分类
                header_len = max(200, len(full_text) // 5)
                return full_text[:header_len].strip()
        except Exception:
            pass
        return None

    @staticmethod
    def _normalize_ocr_text(text: str) -> str:
        """
        统一标准化 OCR 输出文本

        不同 OCR 引擎返回格式不同（纯文本、HTML、Markdown 等），
        在此统一处理为纯文本，确保下游分类器和文件编号提取的行为一致。
        """
        import re
        if not text:
            return text
        # 去除 HTML 标签（云服务可能返回 <div>、<table> 等）
        text = re.sub(r'<[^>]+>', ' ', text)
        # 去除 Markdown 图片语法 ![alt](url)
        text = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', text)
        # 去除 LaTeX 公式（部分模型会产生无意义公式如 \( \text{EICAT} \)）
        text = re.sub(r'\\\(.*?\\\)', '', text)
        text = re.sub(r'\\\[.*?\\\]', '', text)
        # 统一连字符变体为 ASCII 连字符（OCR 常把 - 识别为中文一、破折号等）
        text = re.sub(r'[—–−ー一]', '-', text)
        # 合并多余空格
        text = re.sub(r' +', ' ', text)
        # 去除空行中的多余空格
        text = re.sub(r'\n +', '\n', text)
        # 合并多个连续空行为一个
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    def _ocr_page(self, page) -> str:
        """
        对页面进行 OCR 识别

        Args:
            page: PyMuPDF 页面对象

        Returns:
            识别的文本
        """
        # 延迟初始化 OCR
        self._init_ocr()

        # 导入必要的模块
        import tempfile
        from PIL import Image

        # 检查页面方向（横向/纵向）
        rect = page.rect
        is_landscape = rect.width > rect.height

        # 将页面转为图片并保存为临时文件
        pix = page.get_pixmap(dpi=self.dpi)

        # 检测内容方向：如果页面是纵向但内容可能是横向的
        # 通过检查文本块的方向来判断
        if not is_landscape:
            # 尝试提取文本块来判断内容方向
            text_blocks = page.get_text("blocks")
            if text_blocks:
                # 计算文本块的平均宽高比
                total_ratio = 0
                count = 0
                for block in text_blocks:
                    if len(block) >= 5:  # 确保有坐标信息
                        x0, y0, x1, y1 = block[:4]
                        width = x1 - x0
                        height = y1 - y0
                        if height > 0:
                            ratio = width / height
                            total_ratio += ratio
                            count += 1

                # 如果平均宽高比 < 0.5，说明内容可能是横向排版（文本块很窄很高）
                if count > 0 and (total_ratio / count) < 0.5:
                    print(f"    检测到纵向页面但内容横向排版，旋转处理...")
                    is_landscape = True  # 标记为需要旋转

        # 如果是横向页面或内容横向排版，旋转90度以便更好地识别
        if is_landscape:
            print(f"    检测到横向页面 ({rect.width:.0f}x{rect.height:.0f})，旋转处理...")

            # 保存原始图片
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_path_orig = tmp_file.name
                pix.save(tmp_path_orig)

            # 使用 PIL 旋转图片
            img = Image.open(tmp_path_orig)
            # 尝试旋转90度（顺时针），使页眉在顶部
            img_rotated = img.rotate(-90, expand=True)

            # 保存旋转后的图片
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_path = tmp_file.name
                img_rotated.save(tmp_path)

            # 清理原始临时文件
            try:
                os.unlink(tmp_path_orig)
            except:
                pass
        else:
            # 纵向页面，直接保存
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_path = tmp_file.name
                pix.save(tmp_path)

        # OCR 识别
        try:
            text = self.ocr.recognize_general(tmp_path)
            return text
        finally:
            # 删除临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass

    def _extract_text_single(self, page_file: str) -> dict:
        """
        提取单页文本（线程池任务）

        Args:
            page_file: 单页 PDF 文件路径

        Returns:
            包含 page_num、page_file、text 的字典
        """
        page_num = int(Path(page_file).stem.split('_')[1])
        try:
            text = self.extract_text_from_pdf(page_file)
        except Exception as e:
            print(f"  ⚠️ 第 {page_num} 页文本提取失败: {e}")
            text = ""
        return {'page_num': page_num, 'page_file': page_file, 'text': text}

    def _extract_texts_parallel(self, page_files: list) -> list:
        """
        并行提取所有页面文本（本地直接加载模式自动切换为串行）

        Args:
            page_files: 单页 PDF 文件路径列表

        Returns:
            按原始顺序排列的文本提取结果列表
        """
        # 预初始化 OCR（避免多线程竞争）
        self._init_ocr()

        total = len(page_files)

        # 检查是否使用本地直接加载模式
        use_serial = False
        if self.ocr_engine == "local":
            if hasattr(self.ocr, 'use_direct') and self.ocr.use_direct:
                use_serial = True
                print("  （本地直接加载模式，使用串行处理以避免崩溃）")

        if use_serial:
            # 串行处理（本地直接加载模式）
            results = []
            for i, page_file in enumerate(page_files):
                try:
                    result = self._extract_text_single(page_file)
                    results.append(result)
                    print(f"  [{i+1}/{total}] 第 {result['page_num']} 页 文本提取完成")
                except Exception as e:
                    page_num = int(Path(page_file).stem.split('_')[1])
                    print(f"  [{i+1}/{total}] 第 {page_num} 页 ⚠️ 提取异常: {e}")
                    results.append({
                        'page_num': page_num,
                        'page_file': page_file,
                        'text': ''
                    })
            return results
        else:
            # 并行处理（其他 OCR 引擎）
            from concurrent.futures import ThreadPoolExecutor, as_completed

            results = [None] * total
            completed_count = 0

            with ThreadPoolExecutor(max_workers=min(self.max_workers, total)) as executor:
                future_to_idx = {
                    executor.submit(self._extract_text_single, pf): i
                    for i, pf in enumerate(page_files)
                }
                for future in as_completed(future_to_idx):
                    idx = future_to_idx[future]
                    completed_count += 1
                    try:
                        result = future.result()
                        results[idx] = result
                        print(f"  [{completed_count}/{total}] 第 {result['page_num']} 页 文本提取完成")
                    except Exception as e:
                        page_num = int(Path(page_files[idx]).stem.split('_')[1])
                        print(f"  [{completed_count}/{total}] 第 {page_num} 页 ⚠️ 提取异常: {e}")
                        results[idx] = {
                            'page_num': page_num,
                            'page_file': page_files[idx],
                            'text': ''
                        }

            return results

    def process_pdf(
        self,
        pdf_path: str,
        project_name: str = "",
        system_name: str = "",
        output_base_dir: str = "."
    ):
        """
        处理 PDF：拆分、识别、重命名、归档

        Args:
            pdf_path: PDF 文件路径
            project_name: 项目名称
            system_name: 系统名称（可选，如果不提供则从Word文件读取）
            output_base_dir: 输出基础目录
        """
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            print(f"错误：文件不存在 {pdf_path}")
            return

        output_base_dir = Path(output_base_dir)
        pdf_dir = pdf_path.parent
        code_conflicts = []  # 收集编号冲突告警

        # 0. 读取系统名称（如果未提供）
        system_names_map = {}
        if not system_name:
            print("\n尝试从项目实施申请单读取系统名称...")
            system_names_map = self.read_system_names_from_word(pdf_dir)

        # 1. 拆分 PDF（使用 PDF 文件名作为临时目录名，避免多进程冲突）
        temp_dir = output_base_dir / f"temp_split_{pdf_path.stem}"
        page_files = self.split_pdf(pdf_path, temp_dir)

        # 2. 并行提取所有页面文本
        print(f"\n开始识别文档类型...（{min(self.max_workers, len(page_files))} 并发）")
        text_results = self._extract_texts_parallel(page_files)

        # 3. 串行处理分类
        results = []

        for tr in text_results:
            page_num = tr['page_num']
            page_file = tr['page_file']
            text = tr['text']
            print(f"\n处理第 {page_num} 页:")
            print(f"  提取文本: {text[:100]}..." if len(text) > 100 else f"  提取文本: {text}")

            # 分类文档
            doc_type, is_system_level = self.classifier.classify(text)

            if doc_type:
                print(f"  识别类型: {doc_type} ({'系统级' if is_system_level else '项目级'})")

                # 确定使用的系统名称
                current_system_name = system_name
                system_name_source = None  # 记录系统名称来源：'file_code', 'fuzzy_match', 'content_extract'
                file_code_system_seq = None  # 记录文件编号中的系统序号

                if is_system_level and not system_name:
                    # 方式1：从文件编号 + Word映射表获取系统名称
                    if system_names_map:
                        file_code = self.classifier.extract_file_code(text)
                        print(f"  提取文件编号: {file_code}")
                        if file_code:
                            try:
                                system_seq = int(file_code['system_code'])
                                file_code_system_seq = system_seq  # 保存文件编号中的系统序号
                                print(f"  系统序号: {system_seq}")
                                print(f"  ���统名称映射: {system_names_map}")
                                if system_seq in system_names_map:
                                    current_system_name = system_names_map[system_seq]
                                    system_name_source = 'file_code'
                                    print(f"  ✓ 使用系统名称: {current_system_name} (序号 {system_seq})")

                                    # 校验：检查页面内容中是否出现其他系统名称
                                    # 防止编号冲突（如两个系统共用 BGFS-01）导致系统名称错误
                                    for seq, name in system_names_map.items():
                                        if seq == system_seq:
                                            continue
                                        name_core = name.replace('系统', '').replace('平台', '').strip()
                                        current_core = current_system_name.replace('系统', '').replace('平台', '').strip()
                                        # 页面内容包含其他系统名称，且不包含当前系统名称 → 编号有误，纠正
                                        if (name in text or name_core in text) and current_core not in text:
                                            full_code = file_code.get('full_code', f'?-?-{system_seq:02d}')
                                            print(f"  ⚠️  编号冲突：文件编号 {full_code} 指向「{current_system_name}」(序号 {system_seq})，")
                                            print(f"      但页面内容实际为「{name}」(序号 {seq})，已自动纠正")
                                            code_conflicts.append({
                                                'page': page_num,
                                                'file_code': full_code,
                                                'code_system': current_system_name,
                                                'actual_system': name,
                                            })
                                            current_system_name = name
                                            file_code_system_seq = seq
                                            break
                                else:
                                    print(f"  ✗ 未找到序号 {system_seq} 对应的系统名称")
                            except (ValueError, KeyError) as e:
                                print(f"  ✗ 转换失败: {e}")

                    # 方式2：从 Word 系统名称列表中匹配 OCR 文本（模糊匹配）
                    if not current_system_name and system_names_map:
                        matched_seq = None
                        for seq, name in system_names_map.items():
                            # 去掉常见后缀进行模糊匹配
                            name_core = name.replace('系统', '').replace('平台', '').strip()
                            if name in text or name_core in text:
                                current_system_name = name
                                matched_seq = seq
                                system_name_source = 'fuzzy_match'
                                print(f"  ✓ 从文本匹配系统名称: {current_system_name} (模糊匹配，序号 {seq})")
                                break

                        # 验证：如果有文件编号，检查匹配的系统序号是否一致
                        if current_system_name and file_code_system_seq is not None and matched_seq != file_code_system_seq:
                            print(f"  ⚠️  警告：系统名称不一致！")
                            print(f"      - 文件编号中的系统序号: {file_code_system_seq}")
                            print(f"      - 匹配到的系统名称: {current_system_name} (序号 {matched_seq})")
                            print(f"      - 建议检查文件编号或系统名称是否正确")
                            code_conflicts.append({
                                'page': page_num,
                                'file_code': f'序号 {file_code_system_seq:02d}',
                                'code_system': system_names_map.get(file_code_system_seq, f'序号{file_code_system_seq}'),
                                'actual_system': current_system_name,
                            })

                    # 方式3：从 OCR 文本内容提取系统名称
                    if not current_system_name:
                        extracted_name = self.classifier.extract_system_name_from_content(text)
                        if extracted_name:
                            current_system_name = extracted_name
                            system_name_source = 'content_extract'
                            print(f"  ✓ 从内容提取系统名称: {current_system_name}")

                            # 验证：如果有文件编号和Word映射表，检查提取的系统名称是否与文件编号一致
                            if file_code_system_seq is not None and system_names_map:
                                expected_name = system_names_map.get(file_code_system_seq)
                                if expected_name and expected_name != current_system_name:
                                    # 检查是否是同一个系统（去掉"系统"/"平台"后缀比较）
                                    extracted_core = current_system_name.replace('系统', '').replace('平台', '').strip()
                                    expected_core = expected_name.replace('系统', '').replace('平台', '').strip()
                                    if extracted_core not in expected_core and expected_core not in extracted_core:
                                        print(f"  ⚠️  警告：系统名称不一致！")
                                        print(f"      - 文件编号中的系统序号: {file_code_system_seq}")
                                        print(f"      - 实施单中对应的系统名称: {expected_name}")
                                        print(f"      - 页面提取的系统名称: {current_system_name}")
                                        print(f"      - 建议检查文件编号或页面内容是否正确")
                                        code_conflicts.append({
                                            'page': page_num,
                                            'file_code': f'序号 {file_code_system_seq:02d}',
                                            'code_system': expected_name,
                                            'actual_system': current_system_name,
                                        })
                        else:
                            print(f"  ✗ 未能提取系统名称")

                # 生成文件名（不使用文件编号，不添加页码）
                filename = self.classifier.generate_filename(
                    doc_type, text, project_name, current_system_name, is_system_level, page_num
                )

                # 获取目标文件夹
                folder = self.classifier.get_folder_name(doc_type)

                # 提取文件编号（用于后续合并判断）
                file_code = self.classifier.extract_file_code(text)

                results.append({
                    'page': page_num,
                    'source': page_file,
                    'doc_type': doc_type,
                    'is_system_level': is_system_level,
                    'system_name': current_system_name,
                    'file_code': file_code,  # 保存文件编号用于合并
                    'filename': filename,
                    'folder': folder,
                    'text_preview': text[:200]
                })
            else:
                print(f"  识别类型: 未识别")
                results.append({
                    'page': page_num,
                    'source': page_file,
                    'doc_type': None,
                    'is_system_level': False,
                    'system_name': "",
                    'file_code': None,
                    'filename': f"未分类-page_{page_num:03d}",
                    'folder': "未分类",
                    'text_preview': text[:200]
                })

        # 3. 编号冲突告警汇总
        if code_conflicts:
            print("\n" + "=" * 60)
            print(f"⚠️  发现 {len(code_conflicts)} 处文件编号与页面内容不一致：")
            print("-" * 60)
            for c in code_conflicts:
                print(f"  第 {c['page']:2d} 页 | 编号 {c['file_code']}")
                print(f"         编号指向: {c['code_system']}")
                print(f"         实际内容: {c['actual_system']}（已自动纠正）")
            print("-" * 60)
            print("提示：源文档中存在编号错误，建议核实原始文件。")
            print("=" * 60)

        # 4. 显示处理结果摘要
        self._print_summary(results, project_name, system_name)

        # 5. 询问是否执行归档
        print("\n" + "=" * 60)
        confirm = input("是否执行文件归档？(y/n): ").strip().lower()

        if confirm == 'y':
            # 5.1 合并同一文档的多页
            print("\n检查是否需要合并页面...")
            merged_results = self._merge_pages(results, temp_dir)

            # 5.2 使用PDF文件所在目录作为搜索基础目录
            pdf_dir = pdf_path.parent
            self._archive_files(merged_results, pdf_dir)
            print("\n归档完成！")
        else:
            print("\n已取消归档。拆分的文件保存在:", temp_dir)

    def _get_merge_key(self, result: dict) -> str:
        """
        生成合并键，用于判断哪些页面应该合并

        优先级：
        1. 如果有文件编号，使用文件编号作为键
        2. 如果没有文件编号，使用文档类型+系统名称作为键

        Args:
            result: 页面处理结果

        Returns:
            合并键字符串
        """
        file_code = result.get('file_code')

        if file_code:
            # 方案A：基于文件编号（最准确）
            project_code = file_code.get('project_code', '')
            doc_type_code = file_code.get('doc_type_code', '')
            system_code = file_code.get('system_code', '')
            return f"{project_code}-{doc_type_code}-{system_code}"
        else:
            # 方案B：基于文档类型+系统名称（备用）
            doc_type = result.get('doc_type', '')
            system_name = result.get('system_name', '')
            is_system_level = result.get('is_system_level', False)

            if not doc_type:
                # 未识别文档：每页独立，不合并
                return f"unknown-page{result.get('page', 0)}"

            if is_system_level and system_name:
                # 系统级文档：使用系统名称+文档类型
                return f"{system_name}-{doc_type}"
            elif is_system_level:
                # 系统级文档但无法确定系统归属：每页独立，不合并
                # 避免不同系统的同类文档被错误合并
                return f"{doc_type}-page{result.get('page', 0)}"
            else:
                # 项目级文档：使用文档类型
                return doc_type

    def _merge_pages(self, results: List[dict], temp_dir: Path) -> List[dict]:
        """
        合并同一文档的多页

        Args:
            results: 页面处理结果列表
            temp_dir: 临时目录

        Returns:
            合并后的结果列表
        """
        # 1. 按合并键分组
        groups = {}
        for result in results:
            merge_key = self._get_merge_key(result)

            if merge_key not in groups:
                groups[merge_key] = []

            groups[merge_key].append(result)

        # 2. 拆分系统级文档中编号相同但系统不同的页面
        #    例如：BGFS-01 用于两个不同系统的报告复审记录
        split_groups = {}
        for merge_key, group in groups.items():
            if len(group) > 1 and any(r.get('is_system_level') for r in group):
                # 检查系统名称是否一致
                system_names = set()
                for r in group:
                    sn = r.get('system_name', '')
                    if sn:
                        system_names.add(sn)
                if len(system_names) > 1:
                    # 不同系统的页面使用了相同编号，按系统拆分
                    print(f"\n检测到编号 {merge_key} 下有多个系统: {system_names}")
                    for r in group:
                        sn = r.get('system_name', '')
                        sub_key = f"{merge_key}_{sn}" if sn else f"{merge_key}_page{r.get('page', 0)}"
                        if sub_key not in split_groups:
                            split_groups[sub_key] = []
                        split_groups[sub_key].append(r)
                    continue
            split_groups[merge_key] = group
        groups = split_groups

        # 3. 处理每个组
        merged_results = []
        merge_count = 0

        for merge_key, group in groups.items():
            if len(group) == 1:
                # 单页，不需要合并
                merged_results.append(group[0])
            else:
                # 多页，需要合并
                print(f"\n发现多页文档: {merge_key}")
                print(f"  页面: {[r['page'] for r in group]}")

                # 检查页面间隔
                page_nums = sorted([r['page'] for r in group])
                max_gap = 5  # 允许的最大页面间隔

                has_large_gap = False
                for i in range(len(page_nums) - 1):
                    gap = page_nums[i + 1] - page_nums[i]
                    if gap > max_gap:
                        print(f"  警告：页面 {page_nums[i]} 和 {page_nums[i + 1]} 间隔较大 (间隔 {gap})")
                        has_large_gap = True

                # 询问是否合并
                if has_large_gap:
                    confirm = input(f"  页面间隔较大，是否仍要合并？(y/n): ").strip().lower()
                    if confirm != 'y':
                        print(f"  跳过合并，保留为单页")
                        merged_results.extend(group)
                        continue

                # 执行合并
                merged_file = self._merge_pdf_files(group, temp_dir)

                if merged_file:
                    # 使用第一页的信息，但更新source
                    merged_result = group[0].copy()
                    merged_result['source'] = merged_file
                    merged_result['page'] = f"{page_nums[0]}-{page_nums[-1]}"  # 显示页面范围
                    merged_results.append(merged_result)
                    merge_count += 1
                    print(f"  ✓ 已合并 {len(group)} 页")
                else:
                    # 合并失败，保留原始页面
                    print(f"  ✗ 合并失败，保留为单页")
                    merged_results.extend(group)

        if merge_count > 0:
            print(f"\n共合并 {merge_count} 个文档")

        return merged_results

    def _merge_pdf_files(self, pages: List[dict], temp_dir: Path) -> Optional[str]:
        """
        合并多个PDF文件

        Args:
            pages: 要合并的页面列表
            temp_dir: 临时目录

        Returns:
            合并后的文件路径，失败返回None
        """
        try:
            # 按页码排序
            pages = sorted(pages, key=lambda p: p['page'])

            # 创建新的PDF
            merged_pdf = fitz.open()

            # 逐页添加
            for page in pages:
                src_pdf = fitz.open(page['source'])
                merged_pdf.insert_pdf(src_pdf)
                src_pdf.close()

            # 生成合并后的文件名
            first_page = pages[0]['page']
            last_page = pages[-1]['page']
            merged_filename = f"merged_{first_page}_{last_page}.pdf"
            merged_path = temp_dir / merged_filename

            # 保存
            merged_pdf.save(str(merged_path))
            merged_pdf.close()

            return str(merged_path)

        except Exception as e:
            print(f"  合并失败: {e}")
            return None

    def _print_summary(self, results: List[dict], project_name: str, system_name: str):
        """打印处理结果摘要"""
        print("\n" + "=" * 60)
        print("处理结果摘要")
        print("=" * 60)
        print(f"项目名称: {project_name or '(未指定)'}")
        print(f"系统名称: {system_name or '(未指定)'}")
        print(f"总页数: {len(results)}")
        print()

        # 按文档类型分组统计
        type_counts = {}
        for r in results:
            doc_type = r['doc_type'] or '未识别'
            type_counts[doc_type] = type_counts.get(doc_type, 0) + 1

        print("文档类型统计:")
        for doc_type, count in sorted(type_counts.items()):
            print(f"  {doc_type}: {count} 页")

        print("\n详细列表:")
        for r in results:
            print(f"  第 {r['page']:2d} 页 → {r['filename']}.pdf")
            print(f"           类型: {r['doc_type'] or '未识别'}")
            print(f"           文件夹: {r['folder']}")
            print()

    def _find_matching_folder(self, base_dir: Path, folder_keyword: str) -> Optional[Path]:
        """
        在基础目录下搜索包含关键词的文件夹

        Args:
            base_dir: 基础目录（PDF文件所在目录）
            folder_keyword: 文件夹关键词（如：现场接收归还文档清单）

        Returns:
            匹配的文件夹路径，如果未找到则返回None
        """
        if not base_dir.exists():
            return None

        # 收集所有匹配的文件夹
        candidates = []
        for item in base_dir.iterdir():
            if item.is_dir() and folder_keyword in item.name:
                candidates.append(item)

        if not candidates:
            return None

        # 选择名称最短的文件夹（最精确匹配）
        return min(candidates, key=lambda x: len(x.name))

    def _archive_files(self, results: List[dict], pdf_dir: Path):
        """
        归档文件到对应文件夹

        在PDF文件所在目录下搜索匹配的文件夹，而不是创建新文件夹

        Args:
            results: 处理结果列表
            pdf_dir: PDF文件所在目录
        """
        print("\n开始归档文件...")
        print(f"搜索目录: {pdf_dir}")

        for r in results:
            folder_keyword = r['folder']

            if folder_keyword is None:
                print(f"  警告：'{r.get('type', '未知')}' 没有对应的文件夹映射，归入未分类")
                folder_keyword = "未分类"

            # 搜索匹配的文件夹
            target_folder = self._find_matching_folder(pdf_dir, folder_keyword)

            if target_folder is None:
                # 如果是"未分类"文件夹，自动创建
                if folder_keyword == "未分类":
                    target_folder = pdf_dir / "未分类"
                    target_folder.mkdir(exist_ok=True)
                    print(f"  创建文件夹: {target_folder.name}")
                else:
                    print(f"  警告：未找到包含 '{folder_keyword}' 的文件夹，跳过 {r['filename']}.pdf")
                    continue
            else:
                print(f"  找到文件夹: {target_folder.name}")

            # 目标文件路径
            target_file = target_folder / f"{r['filename']}.pdf"

            # 处理文件名冲突
            counter = 1
            while target_file.exists():
                target_file = target_folder / f"{r['filename']}_{counter}.pdf"
                counter += 1

            # 移动文件
            source = Path(r['source'])
            try:
                source.rename(target_file)
                print(f"  ✓ {source.name} → {target_folder.name}/{target_file.name}")
            except Exception as e:
                print(f"  ✗ 移动失败: {source.name} - {e}")

    # ==================== 多项目批量归档 ====================

    def process_batch_pdf(
        self,
        pdf_path: str,
        projects_base_dir: str
    ):
        """
        处理包含多个项目文档的 PDF，自动归档到各项目文件夹

        Args:
            pdf_path: PDF 文件路径
            projects_base_dir: 项目根目录（包含所有 JDB 项目文件夹的目录）
        """
        pdf_path = Path(pdf_path)
        projects_base_dir = Path(projects_base_dir)

        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")
        if not projects_base_dir.exists():
            raise FileNotFoundError(f"项目根目录不存在: {projects_base_dir}")

        # 初始化日志
        log_dir = pdf_path.parent / "logs"
        log = Logger(log_dir=str(log_dir), max_display_lines=15)

        try:
            self._do_process_batch(pdf_path, projects_base_dir, log)
        finally:
            log.print_static(f"\n日志文件: {log.log_file_path}")
            log.close()

    def _do_process_batch(self, pdf_path: Path, projects_base_dir: Path, log: Logger):
        """批量处理的核心逻辑"""

        log.section(f"批量归档: {pdf_path.name}")
        log.info(f"项目根目录: {projects_base_dir}")

        # 1. 拆分 PDF
        temp_dir = pdf_path.parent / f"temp_split_{pdf_path.stem}"
        log.info(f"正在拆分 PDF（共 {self._get_pdf_page_count(pdf_path)} 页）...")
        page_files = self.split_pdf(str(pdf_path), str(temp_dir))
        total = len(page_files)
        log.info(f"拆分完成: {total} 页")

        # 2. 并行提取文本
        log.section("提取文本与分类")
        text_results = self._extract_texts_parallel_with_log(page_files, log)

        # 3. 分类并提取项目编号
        results = []
        unrecognized = []

        for tr in text_results:
            page_num = tr['page_num']
            text = tr['text']

            # 分类文档
            doc_type, is_system_level = self.classifier.classify(text)

            # 提取文件编号
            file_code = self.classifier.extract_file_code(text)

            if not doc_type:
                log.warn(f"第 {page_num} 页: 文档类型未识别")
                unrecognized.append({
                    'page': page_num,
                    'source': tr['page_file'],
                    'reason': '文档类型未识别',
                    'unarchived_name': f"第{page_num}页_未识别",
                })
                continue

            if not file_code:
                log.warn(f"第 {page_num} 页: 未找到文件编号 ({doc_type})")
                unrecognized.append({
                    'page': page_num,
                    'source': tr['page_file'],
                    'doc_type': doc_type,
                    'reason': '未找到文件编号',
                    'unarchived_name': f"第{page_num}页_{doc_type}",
                })
                continue

            project_code = file_code['project_code']

            # 搜索项目文件夹
            project_dir = self._find_project_folder(projects_base_dir, project_code)
            if not project_dir:
                log.warn(f"第 {page_num} 页: 未找到项目 {project_code} 的文件夹")
                # 生成有意义的文件名，便于人工处理
                unarchived_name = f"{project_code}-{doc_type}"
                if is_system_level and file_code.get('system_code'):
                    unarchived_name = f"{project_code}-{file_code['doc_type_code']}-{file_code['system_code']}"
                unrecognized.append({
                    'page': page_num,
                    'source': tr['page_file'],
                    'doc_type': doc_type,
                    'project_code': project_code,
                    'file_code': file_code,
                    'reason': f'未找到项目文件夹 {project_code}',
                    'unarchived_name': unarchived_name,
                })
                continue

            # 确定系统名称（系统级文档）
            system_name = ""
            if is_system_level:
                system_name = self._resolve_system_name_for_batch(
                    text, file_code, project_dir, log, page_num
                )

            # 生成文件名
            filename = self.classifier.generate_filename(
                doc_type, text, "", system_name, is_system_level, page_num
            )

            # 获取目标文件夹关键词
            folder_keyword = self.classifier.get_folder_name(doc_type)

            log.detail(f"第 {page_num} 页: {project_code} / {doc_type} → {folder_keyword}")
            log.progress(len(results) + len(unrecognized) + 1, total, f"第 {page_num} 页 → {project_code}")

            results.append({
                'page': page_num,
                'source': tr['page_file'],
                'doc_type': doc_type,
                'is_system_level': is_system_level,
                'system_name': system_name,
                'file_code': file_code,
                'filename': filename,
                'folder_keyword': folder_keyword,
                'project_code': project_code,
                'project_dir': project_dir,
            })

        # 4. 显示摘要
        self._print_batch_summary(results, unrecognized, log)

        # 5. 确认归档
        confirm = log.input("\n是否执行批量归档？(y/n): ").strip().lower()

        if confirm == 'y':
            # 5.1 合并同一文档的多页
            log.section("合并与归档")
            merged_results = self._merge_batch_pages(results, temp_dir, log)

            # 5.2 执行归档
            self._archive_batch_files(merged_results, unrecognized, pdf_path.parent, log)

            log.print_static(f"\n批量归档完成！成功 {len(merged_results)} 份，未归档 {len(unrecognized)} 份")
        else:
            log.print_static(f"\n已取消归档。拆分的文件保存在: {temp_dir}")

    def _get_pdf_page_count(self, pdf_path: Path) -> int:
        """获取 PDF 页数"""
        doc = fitz.open(pdf_path)
        count = len(doc)
        doc.close()
        return count

    def _extract_texts_parallel_with_log(self, page_files: list, log: Logger) -> list:
        """带日志的并行文本提取"""
        total = len(page_files)
        log.info(f"开始提取文本（{min(self.max_workers, total)} 并发）...")

        # 检查是否使用串行模式
        use_serial = False
        if self.ocr_engine == "local":
            if hasattr(self.ocr, 'use_direct') and self.ocr.use_direct:
                use_serial = True

        if use_serial:
            results = []
            for i, page_file in enumerate(page_files):
                try:
                    result = self._extract_text_single(page_file)
                    results.append(result)
                except Exception as e:
                    page_num = int(Path(page_file).stem.split('_')[1])
                    results.append({'page_num': page_num, 'page_file': page_file, 'text': ''})
                log.progress(i + 1, total, f"第 {results[-1]['page_num']} 页")
            return results
        else:
            from concurrent.futures import ThreadPoolExecutor, as_completed
            results = [None] * total
            completed = 0

            with ThreadPoolExecutor(max_workers=min(self.max_workers, total)) as executor:
                future_to_idx = {
                    executor.submit(self._extract_text_single, pf): i
                    for i, pf in enumerate(page_files)
                }
                for future in as_completed(future_to_idx):
                    idx = future_to_idx[future]
                    completed += 1
                    try:
                        result = future.result()
                        results[idx] = result
                        log.progress(completed, total, f"第 {result['page_num']} 页")
                    except Exception as e:
                        page_num = int(Path(page_files[idx]).stem.split('_')[1])
                        results[idx] = {'page_num': page_num, 'page_file': page_files[idx], 'text': ''}
                        log.progress(completed, total, f"第 {page_num} 页（异常）")

            return results

    def _find_project_folder(self, base_dir: Path, project_code: str) -> Optional[Path]:
        """
        根据项目编号查找项目文件夹

        Args:
            base_dir: 项目根目录
            project_code: 项目编号（如 JDB25300）

        Returns:
            项目文件夹路径，未找到返回 None
        """
        if not base_dir.exists():
            return None

        for item in base_dir.iterdir():
            if item.is_dir() and item.name.startswith(project_code):
                return item

        return None

    def _resolve_system_name_for_batch(
        self, text: str, file_code: dict, project_dir: Path,
        log: Logger, page_num: int
    ) -> str:
        """
        批量模式下解析系统名称

        优先级：
        1. 从实施单系统名称列表在OCR文本中匹配
        2. 从文本内容正则提取（降级方案）
        """
        system_name = ""

        # 从项目实施单获取系统名称映射（带缓存）
        if not hasattr(self, '_batch_system_names_cache'):
            self._batch_system_names_cache = {}

        project_code = file_code['project_code']
        if project_code not in self._batch_system_names_cache:
            system_names_map = self.read_system_names_from_word(project_dir)
            self._batch_system_names_cache[project_code] = system_names_map
            if not system_names_map:
                log.warn(f"项目 {project_code} 未找到实施单或读取失败，系统级文档将从OCR文本提取系统名称（可能不准确）")
        else:
            system_names_map = self._batch_system_names_cache[project_code]

        # 方式1：用实施单的系统名称在OCR文本中匹配（不依赖OCR识别的序号）
        if system_names_map:
            matched_name = None
            matched_seq = None
            for seq, name in system_names_map.items():
                # 完整名称匹配
                if name in text:
                    matched_name = name
                    matched_seq = seq
                    break
                # 去掉常见后缀进行模糊匹配
                name_core = name.replace('系统', '').replace('平台', '').replace('网络', '').strip()
                if name_core and name_core in text:
                    matched_name = name
                    matched_seq = seq
                    break

            if matched_name:
                system_name = matched_name
                log.detail(f"  第 {page_num} 页: 系统名称={system_name} (实施单匹配，序号 {matched_seq})")

                # 交叉验证：文件编号中的序号 vs 实施单匹配到的序号
                if file_code and file_code.get('system_code'):
                    try:
                        code_seq = int(file_code['system_code'])
                        if code_seq != matched_seq:
                            log.warn(
                                f"第 {page_num} 页: 序号不一致！"
                                f"文件编号 {file_code.get('full_code', '')} 中的系统序号={code_seq:02d}，"
                                f"但实施单匹配到的系统名称「{matched_name}」序号={matched_seq:02d}。"
                                f"请人工确认文件编号或实施单是否正确"
                            )
                        else:
                            log.detail(f"  第 {page_num} 页: 序号交叉验证通过 (文件编号序号={code_seq:02d}, 实施单序号={matched_seq:02d})")
                    except (ValueError, TypeError):
                        pass
            else:
                # 实施单有数据但文本中匹配不到
                names_list = ', '.join(f"{seq}:{name}" for seq, name in system_names_map.items())
                log.warn(f"第 {page_num} 页: OCR文本中未匹配到实施单中的任何系统名称 [{names_list}]")

        # 方式2：从文本内容正则提取（降级方案，准确性较低）
        if not system_name:
            extracted = self.classifier.extract_system_name_from_content(text)
            if extracted:
                system_name = extracted
                log.warn(f"第 {page_num} 页: 系统名称从OCR文本提取: {system_name}（未经实施单验证，请人工确认）")

                # 交叉验证：正则提取的名称与实施单比对
                if system_names_map:
                    # 检查提取的名称是否在实施单中
                    found_in_map = False
                    extracted_core = extracted.replace('系统', '').replace('平台', '').replace('网络', '').strip()
                    for seq, name in system_names_map.items():
                        name_core = name.replace('系统', '').replace('平台', '').replace('网络', '').strip()
                        if extracted == name or (extracted_core and extracted_core == name_core):
                            found_in_map = True
                            # 还可以验证序号
                            if file_code and file_code.get('system_code'):
                                try:
                                    code_seq = int(file_code['system_code'])
                                    if code_seq != seq:
                                        log.warn(
                                            f"第 {page_num} 页: 序号不一致！"
                                            f"文件编号序号={code_seq:02d}，"
                                            f"实施单中「{name}」序号={seq:02d}。"
                                            f"请人工确认"
                                        )
                                except (ValueError, TypeError):
                                    pass
                            break
                    if not found_in_map:
                        names_list = ', '.join(f"{seq}:{name}" for seq, name in system_names_map.items())
                        log.warn(
                            f"第 {page_num} 页: 正则提取的系统名称「{extracted}」不在实施单中 [{names_list}]，请人工确认"
                        )
            else:
                log.warn(f"第 {page_num} 页: 无法确定系统名称，系统级文档将使用默认命名")

        return system_name

    def _print_batch_summary(self, results: list, unrecognized: list, log: Logger):
        """打印批量处理摘要"""
        log.section("处理结果摘要")

        # 按项目分组统计
        project_stats = {}
        for r in results:
            pc = r['project_code']
            if pc not in project_stats:
                project_stats[pc] = {'dir': r['project_dir'].name, 'docs': []}
            project_stats[pc]['docs'].append(r)

        lines = []
        lines.append(f"识别成功: {len(results)} 页，涉及 {len(project_stats)} 个项目")
        lines.append(f"未能归档: {len(unrecognized)} 页")
        lines.append("")

        for pc, info in sorted(project_stats.items()):
            lines.append(f"  {pc} ({info['dir']})")
            type_counts = {}
            for doc in info['docs']:
                dt = doc['doc_type']
                type_counts[dt] = type_counts.get(dt, 0) + 1
            for dt, cnt in sorted(type_counts.items()):
                lines.append(f"    {dt}: {cnt} 页")

        if unrecognized:
            lines.append("")
            lines.append("  未归档的页面:")
            for u in unrecognized:
                name_hint = u.get('unarchived_name', '')
                if name_hint:
                    lines.append(f"    第 {u['page']} 页: {u['reason']} → {name_hint}.pdf")
                else:
                    lines.append(f"    第 {u['page']} 页: {u['reason']}")

        # 打印为静态输出
        for line in lines:
            log.print_static(line)

    def _merge_batch_pages(self, results: list, temp_dir: Path, log: Logger) -> list:
        """
        批量模式的页面合并

        按 项目编号+文件编号 分组合并
        """
        groups = {}
        for r in results:
            merge_key = self._get_merge_key(r)
            project_key = f"{r['project_code']}_{merge_key}"
            if project_key not in groups:
                groups[project_key] = []
            groups[project_key].append(r)

        merged = []
        for key, group in groups.items():
            if len(group) == 1:
                merged.append(group[0])
            else:
                # 按页码排序
                group.sort(key=lambda x: x['page'])
                pages = [g['page'] for g in group]
                log.info(f"合并多页文档: {key}（页面 {pages}）")

                # 检查页面间隔是否合理（最大间隔5页）
                max_gap = 0
                for i in range(1, len(pages)):
                    gap = pages[i] - pages[i - 1]
                    max_gap = max(max_gap, gap)

                if max_gap > 5:
                    log.warn(f"  页面间隔过大（最大 {max_gap}），不合并")
                    merged.extend(group)
                    continue

                # 合并 PDF 文件
                merged_file = self._merge_pdf_files(group, temp_dir)
                if merged_file:
                    merged_result = group[0].copy()
                    merged_result['source'] = merged_file
                    merged_result['page'] = pages[0]
                    merged.append(merged_result)
                    log.info(f"  合并完成: {len(group)} 页 → 1 份")
                else:
                    merged.extend(group)

        return merged

    def _archive_batch_files(self, results: list, unrecognized: list, pdf_dir: Path, log: Logger):
        """将文件归档到各自的项目文件夹"""
        log.info(f"开始归档 {len(results)} 份文件...")

        success_count = 0
        fail_count = 0

        for i, r in enumerate(results):
            project_dir = r['project_dir']
            folder_keyword = r['folder_keyword']

            # 搜索目标子文件夹
            target_folder = self._find_matching_folder(project_dir, folder_keyword) if folder_keyword else None

            if not target_folder:
                log.warn(f"项目 {r['project_code']} 未找到子文件夹 '{folder_keyword}'")
                # 放入未归档文件夹
                unrecognized.append({
                    'page': r['page'],
                    'source': r['source'],
                    'doc_type': r.get('doc_type', ''),
                    'project_code': r['project_code'],
                    'reason': f"项目中未找到子文件夹 '{folder_keyword}'",
                    'unarchived_name': f"{r['project_code']}-{r['filename']}",
                })
                fail_count += 1
                continue

            # 目标文件路径
            target_file = target_folder / f"{r['filename']}.pdf"

            # 处理文件名冲突
            counter = 1
            while target_file.exists():
                target_file = target_folder / f"{r['filename']}_{counter}.pdf"
                counter += 1

            # 移动文件
            source = Path(r['source'])
            try:
                source.rename(target_file)
                success_count += 1
                log.detail(f"{source.name} → {r['project_code']}/{target_folder.name}/{target_file.name}")
                log.progress(i + 1, len(results), f"{r['project_code']}/{target_file.name}")
            except Exception as e:
                log.error(f"移动失败: {source.name} - {e}")
                fail_count += 1

        # 处理未归档的文件
        if unrecognized:
            unarchived_dir = pdf_dir / "未归档"
            unarchived_dir.mkdir(exist_ok=True)
            log.info(f"将 {len(unrecognized)} 份未归档文件移动到: {unarchived_dir.name}/")

            for u in unrecognized:
                source = Path(u['source'])
                if source.exists():
                    # 使用有意义的文件名：项目编号-文档类型.pdf
                    unarchived_name = u.get('unarchived_name', source.stem)
                    target = unarchived_dir / f"{unarchived_name}.pdf"
                    counter = 1
                    while target.exists():
                        target = unarchived_dir / f"{unarchived_name}_{counter}.pdf"
                        counter += 1
                    try:
                        source.rename(target)
                        log.detail(f"  未归档: {source.name} → 未归档/{target.name} ({u['reason']})")
                    except Exception as e:
                        log.error(f"  移动未归档文件失败: {source.name} - {e}")

        log.info(f"归档统计: 成功 {success_count}，失败 {fail_count}，未归档 {len(unrecognized)}")


def main():
    """主函数 - 交互式命令行界面"""
    import argparse

    # 解析命令行参数
    parser = argparse.ArgumentParser(
        description='PDF 文档自动拆分和归档工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
OCR 引擎选项:
  local       - 本地 PaddleOCR-VL（免费，隐私保护）
  aliyun      - 阿里云 OCR（快速稳定，付费）
  siliconflow - SiliconFlow API（免费额度）
  baidu       - 百度飞桨 PaddleOCR 云服务（免费额度）

示例:
  python main.py --ocr local
  python main.py --ocr aliyun
  python main.py --ocr siliconflow
  python main.py --ocr baidu
        """
    )
    parser.add_argument(
        '--ocr',
        choices=['local', 'aliyun', 'siliconflow', 'baidu'],
        help='指定 OCR 引擎'
    )
    parser.add_argument(
        '--batch',
        action='store_true',
        help='批量模式：处理包含多个项目文档的 PDF，自动归档到各项目文件夹'
    )
    parser.add_argument(
        '--projects-dir',
        type=str,
        help='批量模式下的项目根目录路径'
    )

    args = parser.parse_args()

    # 自动检测或使用指定的 OCR 引擎
    ocr_engine = "aliyun"  # 默认
    ocr_display_name = "阿里云 OCR（快速稳定）"

    # 如果命令行指定了引擎，优先使用
    if args.ocr:
        ocr_engine = args.ocr if args.ocr != 'siliconflow' else 'deepseek'
        engine_names = {
            'local': 'PaddleOCR-VL-1.5（本地）',
            'aliyun': '阿里云 OCR（快速稳定）',
            'siliconflow': 'PaddleOCR-VL-1.5（SiliconFlow API）',
            'baidu': '百度飞桨 PaddleOCR（云服务，免费额度）'
        }
        ocr_display_name = engine_names.get(args.ocr, ocr_display_name)
    else:
        # 自动检测
        # 优先检测直接加载模式
        if os.getenv('LOCAL_PADDLEOCR_DIRECT', '').lower() == 'true':
            ocr_engine = "local"
            model_path = os.getenv('LOCAL_PADDLEOCR_MODEL', 'mlx-community/PaddleOCR-VL-1.5-bf16')
            ocr_display_name = f"PaddleOCR-VL-1.5（本地直接加载）"
        else:
            # 检测本地 PaddleOCR-VL HTTP 服务
            local_url = os.getenv('LOCAL_PADDLEOCR_URL', 'http://localhost:8111')
            local_available = False
            try:
                import requests
                response = requests.get(f"{local_url}/health", timeout=2)
                if response.status_code == 200:
                    local_available = True
            except:
                pass

            if local_available:
                ocr_engine = "local"
                ocr_display_name = "PaddleOCR-VL-1.5（本地 MLX-VLM 服务）"
            elif os.getenv('ALIYUN_OCR_APPCODE'):
                ocr_engine = "aliyun"
                ocr_display_name = "阿里云 OCR（快速稳定）"
            elif os.getenv('SILICONFLOW_API_KEY'):
                ocr_engine = "deepseek"
                ocr_display_name = "PaddleOCR-VL-1.5（专业OCR模型）"
            elif os.getenv('BAIDU_PADDLEOCR_TOKEN'):
                ocr_engine = "baidu"
                ocr_display_name = "百度飞桨 PaddleOCR-VL-1.5（云服务，免费额度）"

    print("=" * 60)
    if args.batch:
        print("PDF 多项目批量归档工具")
    else:
        print("PDF 文档自动拆分和重命名工具")
    print(f"使用 {ocr_display_name}")
    print("=" * 60)

    # 输入 PDF 路径
    pdf_path = input("\n请输入 PDF 文件路径: ").strip().strip("'\"")
    if not pdf_path:
        print("错误：未输入文件路径")
        return

    # 创建处理器（本地 OCR 默认并发 1，云端默认并发 3，百度飞桨默认并发 2 避免限流）
    if ocr_engine == "local":
        workers = 1
    elif ocr_engine == "baidu":
        workers = 2
    else:
        workers = 3
    try:
        processor = PDFProcessor(ocr_engine=ocr_engine, max_workers=workers)
    except ValueError as e:
        print(f"\n错误: {e}")
        sys.exit(1)

    if args.batch:
        # 批量模式
        projects_dir = args.projects_dir
        if not projects_dir:
            projects_dir = input("请输入项目根目录路径（包含所有 JDB 项目文件夹）: ").strip().strip("'\"")
        if not projects_dir:
            print("错误：未输入项目根目录路径")
            return

        try:
            processor.process_batch_pdf(pdf_path, projects_dir)
        except (FileNotFoundError, ValueError) as e:
            print(f"\n错误: {e}")
            sys.exit(1)
    else:
        # 单项目模式
        project_name = input("请输入项目名称（可选）: ").strip()
        system_name = input("请输入系统名称（可选）: ").strip()
        output_dir = input("请输入输出目录（默认为当前目录）: ").strip() or "."

        try:
            processor.process_pdf(pdf_path, project_name, system_name, output_dir)
        except ValueError as e:
            print(f"\n错误: {e}")
            sys.exit(1)


if __name__ == "__main__":
    main()
