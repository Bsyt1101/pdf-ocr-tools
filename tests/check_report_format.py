#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查测评报告的前30段内容
"""

import sys
from pathlib import Path

# 添加 src 目录到 Python 路径
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

try:
    from docx import Document
except ImportError:
    print("错误：未安装 python-docx")
    sys.exit(1)


def check_report_content():
    """检查测评报告内容"""

    report_file = Path("/Users/pr0xy/Pr0xyDrive/中科实数工作资料（公司笔记本）/pdf/JDB25027-广州数语科技有限公司-验收阶段归档/15测评报告（Word+PDF）/05.网络安全等级保护测评报告-2025-广州数语科技有限公司-广州数语旅橙团文旅信息共享服务云平台【终稿】.docx")

    if not report_file.exists():
        print(f"错误：文件不存在: {report_file}")
        return

    print(f"读取文件: {report_file.name}\n")
    print("=" * 80)

    doc = Document(report_file)

    print(f"文档总段落数: {len(doc.paragraphs)}\n")
    print("前30段内容:")
    print("-" * 80)

    for i, para in enumerate(doc.paragraphs[:30], 1):
        text = para.text.strip()
        if text:  # 只显示非空段落
            print(f"[{i:2d}] {text}")

    print("-" * 80)


if __name__ == "__main__":
    check_report_content()
